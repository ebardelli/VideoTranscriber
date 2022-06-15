import wave, math, contextlib, os, sys, json, subprocess
import numpy as np

from vosk import Model, SpkModel, KaldiRecognizer, SetLogLevel

from pathlib import Path
from docx import Document
from docx.shared import Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

import glob

def convert_time_stamp(timestamp: str) -> str:
    """ Function to help convert timestamps from s to H:M:S """
    delta = datetime.timedelta(seconds=float(timestamp))
    seconds = delta - datetime.timedelta(microseconds=delta.microseconds)
    return str(seconds)

def write_docx(data, filename, **kwargs):
    """ Write a transcript from the .json transcription file. """
    output_filename = Path(filename)

    # Initiate Document
    document = Document()
    # A4 Size
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)
    # Font
    font = document.styles["Normal"].font

    font.name = "Calibri"

    # Document title and intro
    title = f"Transcription of filename"
    document.add_heading(title, level=1)
    # Set thresholds for formatting later
    threshold_for_highlight = 0.80
    # Intro
    document.add_paragraph(
        "Transcription using automatic speech recognition and"
        " the 'tscribe' python package."
    )
    document.add_paragraph(
        datetime.datetime.now().strftime("Document produced on %A %d %B %Y at %X.")
    )
    document.add_paragraph(
        f"Highlighted text has less than {int(threshold_for_highlight * 100)}% confidence."
    )

    table = document.add_table(rows=1, cols=3)
    table.style = document.styles["Light List Accent 1"]
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Start the first row
    row_cells = table.add_row().cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Time"
    hdr_cells[1].text = "Speaker"
    hdr_cells[2].text = "Content"

    # Add words
    for result in data:
        row_cells = table.add_row().cells

        try:
            # Write timestamp
            start_time = result['result'][0]['start']
            row_cells[0].text = convert_time_stamp(start_time)

            # Write speaker
            row_cells[1].text = "No spk data"            

            for word in result['result']:

                # Get the word with the highest confidence
                # Write the word
                run = row_cells[2].paragraphs[0].add_run(" " + word["word"])
                if float(word["conf"]) < threshold_for_highlight:
                    font = run.font
                    font.color.rgb = RGBColor(100, 149, 237)
        except KeyError:
            pass

    # Formatting transcript table widthds
    widths = (Inches(0.6), Inches(1), Inches(5.4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Save
    document.save(filename)

def resample_ffmpeg(infile):
    stream = subprocess.Popen(
        ['ffmpeg', '-nostdin', '-loglevel', 'quiet', '-i',
        infile,
        '-ar', '16000','-ac', '1', '-f', 's16le', '-'],
        stdout=subprocess.PIPE)
    return stream

def transcribe_stream(rec, stream):
    result = []

    while True:
        data = stream.stdout.read(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            result.append(json.loads(rec.Result()))
    result.append(json.loads(rec.FinalResult()))
    return result

def format_result(result):
    # Recognize speakers
    speakers = []
    for turn in result:
        current_speaker = turn['spk']
        found = 0
        for speaker in speakers:
            if (cosine_dist(current_speaker, speaker) < 0.1):
                found = 1
                break
        if found == 0:
            speakers.append(current_speaker)


    return final_result

def transcribe(video = "", work_dir = "word_dir/", filename="test"):

    audio_file_name = work_dir + "Audio/" + filename +".wav"

    # Load model
    #model_path = "Models/vosk-model-en-us-0.22/"
    model_path = "Models/vosk-model-en-us-0.22-lgraph/"
    #model_path = "Models/vosk-model-small-en-us-0.15/"
    model = Model(model_path)

    spk_model_path = "Models/vosk-model-spk-0.4"
    spk_model = SpkModel(spk_model_path)

    # Set up recognizer
    rec = KaldiRecognizer(model, 16000)
    rec.SetSpkModel(spk_model)

    rec.SetWords(True)

    result = transcribe_stream(rec, resample_ffmpeg(audio_file_name))

    return result

def save_transcript(result, filename):
    with open("Transcripts/"+filename+".json", "w") as f:
        json.dump(result, f)

    final_result = format_result(result)
    print(final_result)

    with open("Transcripts/"+filename+".txt", "w") as f:
        f.write(final_result)
        f.write("\n")
        f.close()

    write_docx(data, "Transcripts/test.docx")

if __name__ == '__main__':
    # Disable vosk log
    SetLogLevel(-1)

    videos = glob.glob('Videos/*.mp4')
    for video in videos:
        print("Transcribing" + video)

        filename = os.path.splitext(os.path.basename(video))[0]

        result = transcribe(video = video, work_dir = "Videos/", filename = filename)
        save_transcript(result, filename)

